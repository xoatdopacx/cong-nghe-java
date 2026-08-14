import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_report():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Header / Title
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hdr1 = p_hdr.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ ĐÔNG Á\nKHOA CÔNG NGHỆ THÔNG TIN - BỘ MÔN CÔNG NGHỆ PHẦN MỀM\n")
    run_hdr1.font.bold = True
    run_hdr1.font.size = Pt(12)
    run_hdr1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BÁO CÁO THỰC HÀNH LAB 4\n")
    run_title.font.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    run_sub = p_title.add_run("HỌC PHẦN: CÔNG NGHỆ JAVA (IT3242)\n")
    run_sub.font.bold = True
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    run_desc = p_title.add_run("Chủ đề: Xử lý sự kiện, EDT và SwingWorker trong Java Swing\n")
    run_desc.font.italic = True
    run_desc.font.size = Pt(12)

    # Student Info Table
    table_info = doc.add_table(rows=4, cols=2)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_info.autofit = False

    info_data = [
        ("Họ và tên sinh viên:", "Nguyễn Văn Hùng"),
        ("Mã sinh viên:", "20230752"),
        ("Lớp / Học phần:", "Công nghệ Java - IT3242"),
        ("Tên Maven Project:", "lab04-swingworker (Package: vn.edu.eaut.lab4)")
    ]

    for i, (k, v) in enumerate(info_data):
        row = table_info.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.3)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.font.bold = True
        r1.font.size = Pt(11)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(11)
        
        set_cell_background(c1, "F2F4F4")
        set_cell_background(c2, "EAEDED")

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Heading 1 helper
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(11)
        return p

    def add_img(img_path, caption):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(2)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            run_cap = p_cap.add_run(f"Hình minh họa: {caption}")
            run_cap.font.italic = True
            run_cap.font.size = Pt(9.5)
            run_cap.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # NỘI DUNG 1
    add_h1("1. CÂU HỎI CỦNG CỐ VÀ CƠ SỞ LÝ THUYẾT")
    
    q_answers = [
        ("1. EDT (Event Dispatch Thread) trong Java Swing là gì?",
         "Event Dispatch Thread (EDT) là luồng (thread) đặc biệt duy nhất chịu trách nhiệm khởi tạo, vẽ lại giao diện đồ họa (GUI) và xử lý toàn bộ các sự kiện tương tác của người dùng trong Java Swing (như click chuột, phím bấm, resize window). Vì các thành phần Swing không an toàn đa luồng (non-thread-safe), mọi thao tác thay đổi giao diện đều phải diễn ra trên EDT."),
        
        ("2. Vì sao không nên đọc file lớn hoặc truy vấn dữ liệu trực tiếp trong ActionListener?",
         "Phương thức actionPerformed() của ActionListener được thực thi trực tiếp trên EDT. Nếu ta đọc file lớn, truy vấn CSDL hoặc tính toán phức tạp trong phương thức này, luồng EDT sẽ bị nghẽn (block). Kết quả là toàn bộ giao diện người dùng sẽ bị đơ/treo (UI Freeze), ứng dụng không thể phản hồi sự kiện hay vẽ lại cửa sổ, gây ra trải nghiệm rất tệ cho người dùng."),
        
        ("3. Phương thức doInBackground() của SwingWorker dùng để làm gì?",
         "doInBackground() chứa toàn bộ mã nguồn xử lý tác vụ nặng (tính toán lâu, đọc file lớn, truy vấn CSDL). Phương thức này tự động được SwingWorker thực thi trên một background thread (luồng phụ), tách biệt hoàn toàn khỏi EDT, nhờ đó giữ cho giao diện giao diện chính luôn mượt mà và linh hoạt."),
        
        ("4. Phương thức done() được gọi khi nào?",
         "done() tự động được gọi trên EDT ngay sau khi phương thức doInBackground() hoàn thành (dù kết thúc bình thường, gặp ngoại lệ hoặc bị hủy bằng cancel()). Do chạy trên EDT, done() là nơi an toàn tuyệt đối để gọi get() lấy kết quả cuối cùng, cập nhật UI (như ẩn progress bar, mở lại nút bấm) hoặc hiển thị thông báo JOptionPane."),
        
        ("5. Sự khác nhau giữa publish/process và setProgress là gì?",
         "• publish(V... chunks) & process(List<V> chunks): Dùng để truyền dữ liệu trung gian kiểu V từ luồng phụ về EDT trong quá trình tính toán. publish() được gọi trong doInBackground(), và SwingWorker sẽ tự gộp dữ liệu rồi chuyển cho process() chạy an toàn trên EDT.\n• setProgress(int progress): Dùng để cập nhật tiến độ công việc dạng % (từ 0 đến 100). Khi gọi setProgress() trong doInBackground(), nó sẽ kích hoạt sự kiện PropertyChangeEvent 'progress' để JProgressBar cập nhật giao diện."),
        
        ("6. Vì sao Lab 4 là bước chuẩn bị cần thiết trước khi làm Lab 5 tích hợp CSDL?",
         "Trong các ứng dụng thực tế (Lab 5), thao tác truy vấn CSDL (JDBC CRUD) luôn tiềm ẩn độ trễ mạng và đĩa cứng. Kiến thức và kỹ năng làm chủ SwingWorker ở Lab 4 sẽ giúp sinh viên thực thi các câu lệnh SQL, nạp dữ liệu vào JTable bất đồng bộ trong Lab 5 mà không gây treo giao diện.")
    ]

    for q, a in q_answers:
        add_h2(q)
        add_p(a)

    # NỘI DUNG 2
    add_h1("2. KẾT QUẢ THỰC HIỆN 5 BÀI TẬP CÓ GỢI Ý CODE (BÀI 1 - BÀI 5)")

    add_h2("Bài 1: Đồng hồ đếm ngược bằng SwingWorker (CountdownFrame.java)")
    add_p("Yêu cầu: Nhập số giây, sử dụng SwingWorker để đếm ngược từng giây và hiển thị lên JLabel. Trong quá trình đếm, nút 'Bắt đầu' bị vô hiệu hóa.")
    add_p("Mã nguồn triển khai chính:")
    add_p("• Sử dụng publish(i) trong vòng lặp doInBackground() để truyền số giây còn lại.\n• Phương thức process(List<Integer> chunks) nhận giá trị và cập nhật lblTime trên EDT.\n• Phương thức done() bật lại btnStart và hiển thị dialog thông báo 'Hoàn thành!'.")
    add_img("screenshots/Bai01_Countdown.png", "Giao diện Bài 1 - Đồng hồ đếm ngược bằng SwingWorker")

    add_h2("Bài 2: Mô phỏng tiến trình tải dữ liệu (ProgressDemoFrame.java)")
    add_p("Yêu cầu: Nút 'Tải dữ liệu', JProgressBar hiển thị % tiến độ từ 0 đến 100% trong 10 giây.")
    add_p("Mã nguồn triển khai chính:")
    add_p("• doInBackground() gọi setProgress(i) qua vòng lặp từ 0 tới 100 với Thread.sleep(1000).\n• Đăng ký PropertyChangeListener lắng nghe thuộc tính 'progress' để cập nhật progressBar.setValue().")
    add_img("screenshots/Bai02_ProgressDemo.png", "Giao diện Bài 2 - Mô phỏng tiến trình tải dữ liệu")

    add_h2("Bài 3: Tính tổng các số nguyên tố nhỏ hơn N (PrimeSumFrame.java)")
    add_p("Yêu cầu: Nhập N, tính tổng các số nguyên tố nhỏ hơn N bằng SwingWorker<Long, Void>.")
    add_p("Mã nguồn triển khai chính:")
    add_p("• Thuật toán kiểm tra số nguyên tố tối ưu isPrime(n).\n• doInBackground() cập nhật setProgress() theo tỷ lệ phần trăm số nguyên đã duyệt.\n• done() thu thập kết quả tính toán bằng phương thức get() và hiển thị lên nhãn kết quả.")
    add_img("screenshots/Bai03_PrimeSum.png", "Giao diện Bài 3 - Tính tổng các số nguyên tố nhỏ hơn N")

    add_h2("Bài 4: Tìm số Fibonacci thứ N bằng memoization (FibonacciFrame.java)")
    add_p("Yêu cầu: Nhập N, tìm số Fibonacci thứ N sử dụng BigInteger tránh tràn số và Map<Integer, BigInteger> ghi nhớ kết quả.")
    add_p("Mã nguồn triển khai chính:")
    add_p("• Thuật toán khử đệ quy bằng Memoization giúp tính các số Fibonacci lớn cực nhanh.\n• JProgressBar được cài đặt chế độ indeterminate (chạy vô tận) trong khi tính toán.")
    add_img("screenshots/Bai04_Fibonacci.png", "Giao diện Bài 4 - Tìm số Fibonacci thứ N với Memoization")

    add_h2("Bài 5: Đọc file lớn và đếm số dòng (FileLineCounterFrame.java)")
    add_p("Yêu cầu: Chọn file qua JFileChooser, dùng SwingWorker đọc file theo từng dòng bằng BufferedReader và cập nhật JProgressBar theo byte đã đọc.")
    add_p("Mã nguồn triển khai chính:")
    add_p("• Sử dụng Files.size() lấy tổng dung lượng file.\n• Tính % tiến độ đọc: progress = readBytes * 100 / totalBytes.")
    add_img("screenshots/Bai05_FileLineCounter.png", "Giao diện Bài 5 - Đọc file lớn và đếm số dòng")

    # NỘI DUNG 3
    add_h1("3. KẾT QUẢ THỰC HIỆN 5 BÀI TẬP TỰ LÀM (BÀI 6 - BÀI 10)")

    add_h2("Bài 6: Bổ sung chức năng hủy tác vụ (CancelableTaskFrame.java)")
    add_p("Mô tả: Nâng cấp tác vụ với nút 'Hủy tác vụ'. Khi bấm 'Hủy', worker.cancel(true) được gọi. Trong doInBackground() chủ động kiểm tra isCancelled() để thoát sớm. Phương thức done() kiểm tra isCancelled() để cập nhật trạng thái 'Đã hủy tác vụ'.")
    add_img("screenshots/Bai06_CancelableTask.png", "Giao diện Bài 6 - Chức năng hủy tác vụ an toàn với cancel()")

    add_h2("Bài 7: Tìm kiếm từ khóa trong file văn bản lớn (KeywordSearchFrame.java)")
    add_p("Mô tả: Chọn file .txt và nhập từ khóa. SwingWorker đọc file theo dòng, tìm kiếm không phân biệt hoa/thường (contains(keyword.toLowerCase())), dùng publish() để đẩy từng dòng trùng khớp lên JTextArea hiển thị ngay lập tức.")
    add_img("screenshots/Bai07_KeywordSearch.png", "Giao diện Bài 7 - Tìm kiếm từ khóa trong file văn bản lớn")

    add_h2("Bài 8: Đọc file CSV điểm sinh viên và thống kê (StudentCsvStatsFrame.java)")
    add_p("Mô tả: Đọc file CSV chứa danh sách sinh viên bất đồng bộ, nạp từng dòng vào JTable qua publish()/process(). Khi hoàn thành, tính điểm trung bình lớp và tìm ra sinh viên có điểm số cao nhất.")
    add_img("screenshots/Bai08_StudentCsvStats.png", "Giao diện Bài 8 - Nạp bảng JTable và thống kê file CSV sinh viên")

    add_h2("Bài 9: Mô phỏng tải danh sách sản phẩm (ProductLoadDemoFrame.java)")
    add_p("Mô tả: Mô phỏng tải dữ liệu sản phẩm bất đồng bộ từ hệ thống vào JTable, có JProgressBar và nhãn trạng thái. Bài tập tạo bước đệm trực tiếp cho việc nạp dữ liệu từ CSDL MySQL/SQL Server ở Lab 5.")
    add_img("screenshots/Bai09_ProductLoadDemo.png", "Giao diện Bài 9 - Mô phỏng tải dữ liệu sản phẩm bất đồng bộ vào JTable")

    add_h2("Bài 10: Mini Project - Quản lý sản phẩm bằng file CSV (ProductManagerFrame.java)")
    add_p("Mô tả: Ứng dụng tổng hợp quản lý sản phẩm hoàn chỉnh bao gồm các chức năng Thêm, Sửa, Xóa trên JTable và Đọc/Lưu file CSV bất đồng bộ với SwingWorker. Đầy đủ các bước kiểm tra hợp lệ dữ liệu nhập (validation).")
    add_img("screenshots/Bai10_ProductManager.png", "Giao diện Bài 10 - Mini Project Quản lý sản phẩm với CSV & SwingWorker")

    add_h2("Giao diện trung tâm Dashboard Launcher (MainDashboardFrame.java)")
    add_p("Mô tả: Khởi chạy tập trung cho phép giáo viên / người chấm điểm dễ dàng trải nghiệm và kiểm tra toàn bộ 10 bài tập Lab 4 chỉ với 1 click.")
    add_img("screenshots/MainDashboard.png", "Giao diện trung tâm Main Dashboard Launcher cho toàn bộ Lab 4")

    # KẾT LUẬN
    add_h1("4. KẾT LUẬN VÀ TỰ ĐÁNH GIÁ")
    add_p("• Đã hoàn thành xuất sắc 100% tất cả các yêu cầu bài tập Lab 4 (5 bài gợi ý code + 5 bài tự làm).\n• Làm chủ cơ chế đa luồng trong Swing, hiểu rõ vai trò của Event Dispatch Thread (EDT) và sử dụng thành thạo SwingWorker (doInBackground, publish, process, setProgress, done, cancel).\n• Mã nguồn chuẩn Maven với cấu trúc package vn.edu.eaut.lab4, mã sạch, giao diện đẹp mắt và không bị treo UI khi thực hiện các tác vụ nặng.")

    output_path = "BaoCao_Lab4_NguyenVanHung_20230752.docx"
    doc.save(output_path)
    print(f"Đã tạo file báo cáo Word thành công: {output_path}")

if __name__ == '__main__':
    create_report()
